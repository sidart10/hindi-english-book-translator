#!/usr/bin/env python3
"""
Main Controller for Hindi Book Translation System
Orchestrates the translation process with batch processing and concurrency control
"""

import asyncio
import json
import os
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import aiohttp
from asyncio import Semaphore
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import our components
from translation_engine import TranslationEngine, TranslationConfig
from document_processor import DocumentProcessor
from cost_meter import CostMeter


class BookTranslationController:
    """Main controller for book translation workflow"""
    
    def __init__(self, config_path: str = "config.json"):
        """
        Initialize controller with configuration
        
        Args:
            config_path: Path to configuration file
        """
        self.config = self.load_config(config_path)
        
        # Initialize components
        self.translation_engine = TranslationEngine(
            TranslationConfig(**self.config.get("translation", {}))
        )
        self.document_processor = DocumentProcessor()
        self.cost_meter = CostMeter(
            monthly_budget=self.config.get("cost", {}).get("monthly_budget", 1000),
            alert_threshold=self.config.get("cost", {}).get("alert_threshold", 0.7)
        )
        
        # Set up concurrency control
        self.semaphore = Semaphore(
            self.config.get("processing", {}).get("quota_limit", 10)
        )
        
        # Processing parameters
        self.batch_size = self.config.get("processing", {}).get("batch_size", 50)
        self.max_retries = self.config.get("processing", {}).get("max_retries", 3)
        
        print(f"Controller initialized with batch_size={self.batch_size}, "
              f"concurrency={self.config.get('processing', {}).get('quota_limit', 10)}")
    
    def load_config(self, config_path: str) -> Dict:
        """Load configuration from file"""
        if Path(config_path).exists():
            with open(config_path, 'r') as f:
                return json.load(f)
        else:
            # Default configuration
            return {
                "translation": {
                    "source_language": "hi",
                    "target_language": "en",
                    "temperature": 0,
                    "max_characters": 9000
                },
                "processing": {
                    "batch_size": 50,
                    "quota_limit": 10,
                    "max_retries": 3
                },
                "cost": {
                    "monthly_budget": 1000,
                    "alert_threshold": 0.7
                },
                "quality": {
                    "min_confidence": 0.85
                }
            }
    
    async def translate_book(self, input_path: str, output_path: str,
                           project_id: Optional[str] = None) -> Dict:
        """
        Main method to translate entire book
        
        Args:
            input_path: Path to input document
            output_path: Path for output DOCX
            project_id: Optional project identifier
            
        Returns:
            Dictionary with translation results and statistics
        """
        if not project_id:
            project_id = f"book_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        print(f"\n📚 Starting translation of '{input_path}'")
        print(f"   Project ID: {project_id}")
        
        results = {
            "project_id": project_id,
            "input_file": input_path,
            "output_file": output_path,
            "start_time": datetime.now().isoformat(),
            "statistics": {
                "total_sentences": 0,
                "translated_sentences": 0,
                "total_characters": 0,
                "total_cost": 0.0,
                "errors": [],
                "quality_warnings": []
            }
        }
        
        try:
            # Step 1: Process document into sentences
            print("\n📖 Extracting text from document...")
            sentences = list(self.document_processor.process_document(input_path))
            total_sentences = len(sentences)
            results["statistics"]["total_sentences"] = total_sentences
            
            if total_sentences == 0:
                raise ValueError("No sentences extracted from document")
            
            print(f"   Extracted {total_sentences} sentences")
            
            # Step 2: Generate segmentation JSONL
            segmentation_file = self._generate_segmentation_jsonl(sentences, project_id)
            print(f"   Generated segmentation file: {segmentation_file}")
            
            # Step 3: Translate in batches
            print(f"\n🔄 Translating in batches of {self.batch_size}...")
            translated_sentences = await self._translate_batches(sentences)
            
            results["statistics"]["translated_sentences"] = len(translated_sentences)
            results["statistics"]["total_cost"] = self.cost_meter.current_cost
            
            # Step 4: Generate output DOCX
            print(f"\n📄 Generating output document...")
            self._generate_output(translated_sentences, output_path)
            
            # Step 5: Final statistics
            results["end_time"] = datetime.now().isoformat()
            
            # Calculate total characters
            total_chars = 0
            for s in sentences:
                if hasattr(s, 'text'):
                    total_chars += len(s.text)
                else:
                    total_chars += len(s.get("text", ""))
            results["statistics"]["total_characters"] = total_chars
            
            print(f"\n✅ Translation complete!")
            print(f"   Total sentences: {total_sentences}")
            print(f"   Total characters: {total_chars:,}")
            print(f"   Total cost: ${self.cost_meter.current_cost:.2f}")
            print(f"   Output saved to: {output_path}")
            
            # Check for cost alerts
            status = self.cost_meter.get_status()
            if status["alerts"]:
                print("\n⚠️  Cost Alerts:")
                for alert in status["alerts"]:
                    print(f"   {alert}")
            
            return results
            
        except Exception as e:
            results["statistics"]["errors"].append(str(e))
            print(f"\n❌ Error during translation: {e}")
            raise
    
    def _generate_segmentation_jsonl(self, sentences: List[Dict], 
                                    project_id: str) -> str:
        """Generate JSONL file with segmented sentences"""
        jsonl_path = f"segmentation_{project_id}.jsonl"
        
        with open(jsonl_path, 'w', encoding='utf-8') as f:
            for sent in sentences:
                # Convert to dict if it's a DocumentSegment object
                if hasattr(sent, 'to_dict'):
                    sent_dict = sent.to_dict()
                else:
                    sent_dict = sent
                f.write(json.dumps(sent_dict, ensure_ascii=False) + '\n')
        
        return jsonl_path
    
    async def _translate_batches(self, sentences: List[Dict]) -> List[Dict]:
        """Translate sentences in batches with concurrency control"""
        translated_sentences = []
        total_batches = (len(sentences) + self.batch_size - 1) // self.batch_size
        
        for batch_idx in range(0, len(sentences), self.batch_size):
            batch = sentences[batch_idx:batch_idx + self.batch_size]
            current_batch_num = (batch_idx // self.batch_size) + 1
            
            print(f"   Batch {current_batch_num}/{total_batches} "
                  f"({len(batch)} sentences)")
            
            # Check cost threshold before translating
            if self.cost_meter.check_threshold(0.7):
                cost_status = self.cost_meter.get_status()
                print(f"   ⚠️  {cost_status['alerts'][0]}")
            
            # Extract text from sentence objects
            batch_texts = []
            for s in batch:
                if hasattr(s, 'text'):
                    batch_texts.append(s.text)
                else:
                    batch_texts.append(s.get("text", ""))
            
            # Translate with rate limiting and retry logic
            translated = await self._translate_with_retry(batch_texts, batch_idx)
            
            # Combine original and translated data
            for sent_obj, trans_result in zip(batch, translated):
                # Convert DocumentSegment to dict if needed
                if hasattr(sent_obj, 'to_dict'):
                    sent_dict = sent_obj.to_dict()
                else:
                    sent_dict = sent_obj
                
                sent_dict["translation"] = trans_result.get("translation", "")
                sent_dict["confidence"] = trans_result.get("confidence", 0.0)
                
                # Basic quality check
                if trans_result.get("confidence", 0) < self.config.get("quality", {}).get("min_confidence", 0.85):
                    sent_dict["quality_warning"] = "Low confidence translation"
                
                translated_sentences.append(sent_dict)
            
            # Update cost tracking
            total_chars = sum(len(t) for t in batch_texts)
            cost_result = self.cost_meter.add_cost(
                total_chars, 
                f"Batch {current_batch_num}"
            )
            
            print(f"      Cost: ${cost_result['cost']:.4f} "
                  f"(Total: ${cost_result['total_cost']:.2f})")
        
        return translated_sentences
    
    async def _translate_with_retry(self, texts: List[str], 
                                   batch_idx: int) -> List[Dict]:
        """Translate with retry logic for rate limiting and errors"""
        for retry in range(self.max_retries):
            try:
                async with self.semaphore:
                    # Call translation engine
                    results = await self.translation_engine.translate_chunk(texts)
                    return results
                    
            except aiohttp.ClientResponseError as e:
                if e.status == 429:  # Rate limit
                    wait_time = 2 ** retry * 10  # Exponential backoff
                    print(f"      Rate limit hit, waiting {wait_time}s...")
                    await asyncio.sleep(wait_time)
                    
                elif e.status >= 500:  # Server error
                    wait_time = 2 ** retry * 5
                    print(f"      Server error, retrying in {wait_time}s...")
                    await asyncio.sleep(wait_time)
                    
                else:
                    raise  # Other errors, don't retry
                    
            except Exception as e:
                if retry == self.max_retries - 1:
                    print(f"      Failed after {self.max_retries} retries: {e}")
                    # Return empty translations for this batch
                    return [{"translation": "", "error": str(e)} for _ in texts]
                else:
                    wait_time = 2 ** retry * 2
                    print(f"      Error: {e}, retrying in {wait_time}s...")
                    await asyncio.sleep(wait_time)
        
        # Should not reach here
        return [{"translation": "", "error": "Max retries exceeded"} for _ in texts]
    
    def _generate_output(self, translated_sentences: List[Dict], 
                        output_path: str):
        """Generate DOCX output with translations"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Hindi to English Translation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Translation Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Total Sentences: {len(translated_sentences)}")
        doc.add_paragraph("")
        
        # Track page numbers
        current_page = None
        
        for sent in translated_sentences:
            # Add page break when page changes
            if sent.get("page") != current_page:
                if current_page is not None:
                    doc.add_page_break()
                current_page = sent.get("page")
                
                # Add page header
                page_header = doc.add_heading(f'Page {current_page}', level=1)
                page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")
            
            # Add translated text
            para = doc.add_paragraph(sent.get("translation", ""))
            
            # Add quality warning if present
            if sent.get("quality_warning"):
                warning = doc.add_paragraph(
                    f"⚠️ {sent['quality_warning']} (confidence: {sent.get('confidence', 0):.2f})"
                )
                warning.style = 'Caption'
                
                # Make warning text smaller and gray
                for run in warning.runs:
                    run.font.size = Pt(10)
        
        # Save document
        doc.save(output_path)
        print(f"   Document saved to: {output_path}")


# Test function
async def test_controller():
    """Test the main controller with a sample text file"""
    print("Testing Main Controller...")
    print("=" * 70)
    
    # Create a test text file
    test_file = "test_controller_input.txt"
    with open(test_file, 'w', encoding='utf-8') as f:
        f.write("नमस्ते, मेरा नाम राम है।\n")
        f.write("मैं भारत से हूं।\n")
        f.write("यह एक परीक्षण है।\n")
        f.write("आज मौसम बहुत अच्छा है।\n")
        f.write("मुझे हिंदी भाषा पसंद है।\n")
    
    # Initialize controller
    controller = BookTranslationController()
    
    # Run translation
    output_file = "test_controller_output.docx"
    
    try:
        results = await controller.translate_book(
            test_file,
            output_file,
            project_id="test_001"
        )
        
        print("\n" + "=" * 70)
        print("Translation Results:")
        print(json.dumps(results, indent=2, ensure_ascii=False))
        
    except Exception as e:
        print(f"Test failed: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # Clean up test files
        if Path(test_file).exists():
            os.remove(test_file)
        if Path(f"segmentation_test_001.jsonl").exists():
            os.remove(f"segmentation_test_001.jsonl")


if __name__ == "__main__":
    # Run test
    asyncio.run(test_controller()) 